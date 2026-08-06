import base64
import hashlib
import html
import json
import os
import re
import tempfile
import urllib.parse
import urllib.request

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import gradio as gr
import markdown as md
from weasyprint import HTML

from themes import THEMES, Theme

EXPORT_DIR = os.path.abspath('exports')
os.makedirs(EXPORT_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Word count logic
# ---------------------------------------------------------------------------

def analyze_text(text: str) -> tuple[int, int, int, int, int, int, int]:
    '''
    Analyze input text and calculate character and word count metrics.

    Args:
        text: The input string pasted or typed by the user.

    Returns:
        A tuple containing:
            - total_chars: Total length including spaces and newlines.
            - chars_no_spaces: Total characters excluding whitespace.
            - cjk_count: Count of CJK Chinese ideographs.
            - cjk_punct_count: Count of Chinese full-width punctuation.
            - english_words: Count of English words.
            - digit_count: Count of numeric digits.
            - line_count: Count of lines.
    '''
    if not text:
        return 0, 0, 0, 0, 0, 0, 0

    total_chars = len(text)
    chars_no_spaces = len(re.sub(r'\s+', '', text))

    # CJK Chinese Character matching
    cjk_pattern = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf\U00020000-\U0002a6df]')
    cjk_count = len(cjk_pattern.findall(text))

    # CJK / Full-width Punctuation matching
    cjk_punct_pattern = re.compile(r'[\u3000-\u303f\uff00-\uffef]')
    cjk_punct_count = len(cjk_punct_pattern.findall(text))

    # English word matching
    english_pattern = re.compile(r'\b[a-zA-Z]+(?:-[a-zA-Z]+)?\b')
    english_words = len(english_pattern.findall(text))

    # Digit matching
    digit_pattern = re.compile(r'\d')
    digit_count = len(digit_pattern.findall(text))

    # Line counting
    line_count = len(text.splitlines()) if text else 0

    return (
        total_chars,
        chars_no_spaces,
        cjk_count,
        cjk_punct_count,
        english_words,
        digit_count,
        line_count,
    )


def clear_input() -> tuple[str, int, int, int, int, int, int, int]:
    '''
    Reset the input text and all statistics counters to initial state.

    Returns:
        A tuple with empty string and 0 for all metrics.
    '''
    return '', 0, 0, 0, 0, 0, 0, 0


# ---------------------------------------------------------------------------
# Markdown preview & Mermaid rendering logic
# ---------------------------------------------------------------------------

_MERMAID_BLOCK_RE = re.compile(
    r'```mermaid\s*\n(.*?)```',
    re.DOTALL,
)

_MD_EXTENSIONS = [
    'tables',
    'fenced_code',
    'codehilite',
    'toc',
    'sane_lists',
    'smarty',
]


def _get_export_filename(md_text: str, ext: str = 'pdf') -> str:
    '''
    Determine default export filename based on H1 header or MD5 content hash.

    Args:
        md_text: Raw markdown text string.
        ext: Target file extension without dot (e.g. 'pdf' or 'docx').

    Returns:
        Clean filename string including extension.
    '''
    h1_match = re.search(r'^\s*#\s+(.+)$', md_text, re.MULTILINE)
    if h1_match:
        raw_title = h1_match.group(1).strip()
        raw_title = re.sub(r'[*_`~]', '', raw_title)
        clean_title = re.sub(r'[\\/:*?"<>|\r\n\t]', '_', raw_title).strip(' ._')
        if clean_title:
            return f'{clean_title}.{ext}'

    text_hash = hashlib.md5(md_text.encode('utf-8')).hexdigest()[:8]
    return f'export_{text_hash}.{ext}'


def _fetch_mermaid_png(code: str, mermaid_theme: str = 'default') -> bytes | None:
    '''
    Fetch rendered PNG bytes for a Mermaid diagram from mermaid.ink API.

    Args:
        code: Mermaid diagram source string.
        mermaid_theme: Mermaid theme identifier ('default' or 'dark').

    Returns:
        PNG image bytes, or None if fetching fails.
    '''
    try:
        payload = json.dumps({'code': code, 'mermaid': {'theme': mermaid_theme}})
        b64_str = base64.urlsafe_b64encode(payload.encode('utf-8')).decode('utf-8')
        url = f'https://mermaid.ink/img/{b64_str}'
        req = urllib.request.Request(
            url,
            headers={'User-Agent': 'Mozilla/5.0'}
        )
        with urllib.request.urlopen(req, timeout=8) as response:
            if response.status == 200:
                return response.read()
    except Exception:
        pass
    return None


def _sanitize_emojis_for_export(text: str) -> str:
    '''
    Replace color emojis with clean HTML / Unicode symbols for PDF export.

    Args:
        text: Markdown or HTML text containing color emojis.

    Returns:
        Sanitized text with standard web symbols.
    '''
    text = text.replace('✅', '<span class="chk">✓</span>')
    text = text.replace('❌', '<span class="crs">✗</span>')
    return text


def _replace_mermaid_blocks(md_text: str) -> tuple[str, bool]:
    '''
    Extract mermaid code blocks and replace them with <div class="mermaid">
    placeholders for client-side rendering in browser preview.

    Args:
        md_text: Raw markdown string.

    Returns:
        Tuple of (modified markdown text, whether mermaid blocks were found).
    '''
    has_mermaid = bool(_MERMAID_BLOCK_RE.search(md_text))

    def _replacer(match: re.Match) -> str:
        code = match.group(1).strip()
        return f'<div class="mermaid">\n{code}\n</div>'

    return _MERMAID_BLOCK_RE.sub(_replacer, md_text), has_mermaid


def _build_html_page(body_html: str, theme: Theme, has_mermaid: bool = False) -> str:
    '''
    Wrap rendered HTML body in a full HTML document with CSS styling
    and optional Mermaid.js script.

    Args:
        body_html: The rendered HTML content.
        theme: Selected Theme instance.
        has_mermaid: Whether to include mermaid.js script tag.

    Returns:
        Complete HTML document string.
    '''
    mermaid_script = ''
    if has_mermaid:
        mermaid_script = (
            '<script src="https://cdn.jsdelivr.net/npm/mermaid@11/dist/'
            'mermaid.min.js"></script>\n'
            f'<script>mermaid.initialize({{startOnLoad:true, theme:"{theme.mermaid_theme}"}});</script>'
        )

    return (
        f'<!DOCTYPE html>\n<html lang="zh-Hant">\n<head>\n'
        f'<meta charset="UTF-8">\n'
        f'<style>{theme.css}</style>\n'
        f'{mermaid_script}\n'
        f'</head>\n<body>\n{body_html}\n</body>\n</html>'
    )


def render_preview(md_text: str, theme_name: str = 'Light') -> str:
    '''
    Convert markdown text to styled HTML for live preview using selected theme.

    Uses an iframe with srcdoc to ensure Mermaid.js scripts execute properly
    since Gradio strips script tags from gr.HTML components.

    Args:
        md_text: Raw markdown string.
        theme_name: Selected theme name ('Light', 'Dark', 'Nord', 'Dracula').

    Returns:
        An iframe HTML tag containing the rendered preview, or a placeholder
        message if input is empty.
    '''
    if not md_text or not md_text.strip():
        return (
            '<div style="color:#94a3b8;padding:2rem;text-align:center;'
            'font-size:1.1em;">請在左側輸入 Markdown 文字以預覽</div>'
        )

    theme = THEMES.get(theme_name, THEMES['Light'])
    processed, has_mermaid = _replace_mermaid_blocks(md_text)
    body_html = md.markdown(processed, extensions=_MD_EXTENSIONS)
    full_html = _build_html_page(body_html, theme, has_mermaid)

    escaped = html.escape(full_html, quote=True)
    return (
        f'<iframe class="preview-iframe" srcdoc="{escaped}" '
        f'style="width:100%;height:850px;min-height:850px;border:1px solid #cbd5e1;'
        f'border-radius:8px;background:{theme.body_bg};" '
        f'sandbox="allow-scripts allow-same-origin">'
        f'</iframe>'
    )


# ---------------------------------------------------------------------------
# Export: PDF
# ---------------------------------------------------------------------------

def _render_md_to_html_for_export(md_text: str, theme: Theme) -> str:
    '''
    Convert markdown to a print-ready HTML document with rendered Mermaid diagrams.

    Args:
        md_text: Raw markdown string.
        theme: Selected Theme instance.

    Returns:
        Complete HTML string suitable for weasyprint PDF generation.
    '''
    def _mermaid_replacer(match: re.Match) -> str:
        code = match.group(1).strip()
        png_data = _fetch_mermaid_png(code, theme.mermaid_theme)
        if png_data:
            b64_img = base64.b64encode(png_data).decode('utf-8')
            return f'<div class="mermaid-img"><img src="data:image/png;base64,{b64_img}" alt="Mermaid Diagram" /></div>'
        return f'<pre class="mermaid-fallback">[Mermaid Diagram]\n{html.escape(code)}</pre>'

    processed = _MERMAID_BLOCK_RE.sub(_mermaid_replacer, md_text)
    body_html = md.markdown(processed, extensions=_MD_EXTENSIONS)
    body_html = _sanitize_emojis_for_export(body_html)
    return _build_html_page(body_html, theme, has_mermaid=False)


def export_pdf(md_text: str, theme_name: str = 'Light') -> str | None:
    '''
    Export markdown text to a PDF file styled with the selected theme.

    Args:
        md_text: Raw markdown string.
        theme_name: Selected theme name.

    Returns:
        Path to the generated PDF file, or None if input is empty.
    '''
    if not md_text or not md_text.strip():
        return None

    filename = _get_export_filename(md_text, ext='pdf')
    theme = THEMES.get(theme_name, THEMES['Light'])
    html_str = _render_md_to_html_for_export(md_text, theme)

    tmp_dir = tempfile.mkdtemp(dir=EXPORT_DIR, prefix='pdf_')
    out_path = os.path.join(tmp_dir, filename)
    HTML(string=html_str).write_pdf(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Export: Word (.docx)
# ---------------------------------------------------------------------------

def _set_cjk_font(run, font_name: str = 'Noto Sans CJK TC', size_pt: int = 11):
    '''
    Set font for a run supporting CJK characters.

    Args:
        run: The docx Run object.
        font_name: Font family name.
        size_pt: Font size in points.
    '''
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _append_node_to_paragraph(
    p,
    node,
    theme: Theme,
    is_bold: bool = False,
    is_italic: bool = False,
    is_code: bool = False,
    code_class: str = '',
) -> None:
    '''
    Recursively parse HTML DOM nodes and append formatted runs to a Word paragraph.

    Applies syntax token colors to Pygments highlighted code blocks.

    Args:
        p: Word paragraph object.
        node: BeautifulSoup HTML node.
        theme: Selected Theme instance.
        is_bold: Whether current context is bold text.
        is_italic: Whether current context is italic text.
        is_code: Whether current context is code text.
        code_class: Pygments span class for code syntax highlighting.
    '''
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = p.add_run(text)
            _set_cjk_font(
                run,
                font_name='Noto Sans Mono CJK TC' if is_code else 'Noto Sans CJK TC',
                size_pt=9.5 if is_code else 11,
            )
            run.bold = is_bold
            run.italic = is_italic

            # Apply syntax highlighting colors in code blocks
            if is_code:
                if code_class in ('k', 'kn', 'kr', 'kd'):
                    run.font.color.rgb = RGBColor(0xd7, 0x3a, 0x49)  # Keyword - Red/Pink
                    run.bold = True
                elif code_class in ('s', 's1', 's2', 'sd'):
                    run.font.color.rgb = RGBColor(0x03, 0x2f, 0x62)  # String - Dark Blue/Green
                elif code_class in ('c', 'c1', 'cm', 'ch'):
                    run.font.color.rgb = RGBColor(0x6a, 0x73, 0x7d)  # Comment - Gray
                    run.italic = True
                elif code_class in ('nf', 'fm'):
                    run.font.color.rgb = RGBColor(0x6f, 0x42, 0xc1)  # Function - Purple
                    run.bold = True
                elif code_class in ('mi', 'mf', 'mh'):
                    run.font.color.rgb = RGBColor(0x00, 0x5c, 0x9e)  # Number - Blue
                else:
                    run.font.color.rgb = theme.docx_colors.code_text
            else:
                if '✓' in text:
                    run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
                    run.bold = True
                elif '✗' in text:
                    run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
                    run.bold = True

        return

    tag = node.name.lower() if node.name else ''
    child_bold = is_bold or tag in ('strong', 'b')
    child_italic = is_italic or tag in ('em', 'i')
    child_code = is_code or tag in ('code', 'kbd', 'samp')
    child_class = node.get('class', [''])[0] if node.name else ''

    for child in node.children:
        _append_node_to_paragraph(
            p,
            child,
            theme,
            is_bold=child_bold,
            is_italic=child_italic,
            is_code=child_code,
            code_class=child_class or code_class,
        )


def export_word(md_text: str, theme_name: str = 'Light') -> str | None:
    '''
    Export markdown text to a Word (.docx) file styled with the selected theme.

    Converts Markdown to HTML DOM elements and maps them into Word document
    structures (headings, formatted text runs, lists, tables, blockquotes,
    Pygments code blocks, and embedded Mermaid diagram PNGs).

    Args:
        md_text: Raw markdown string.
        theme_name: Selected theme name.

    Returns:
        Path to the generated .docx file, or None if input is empty.
    '''
    if not md_text or not md_text.strip():
        return None

    filename = _get_export_filename(md_text, ext='docx')
    theme = THEMES.get(theme_name, THEMES['Light'])
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans CJK TC'
    font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK TC')

    # Handle Mermaid diagrams: render to PNG temp files before HTML parsing
    mermaid_images: dict[str, str] = {}

    def _mermaid_word_replacer(match: re.Match) -> str:
        code = match.group(1).strip()
        png_data = _fetch_mermaid_png(code, theme.mermaid_theme)
        placeholder_id = f'MERMAID_IMG_PLACEHOLDER_{len(mermaid_images)}'
        if png_data:
            tmp_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='mermaid_')
            tmp_img.write(png_data)
            tmp_img.close()
            mermaid_images[placeholder_id] = tmp_img.name
            return f'<p class="mermaid-img-p">{placeholder_id}</p>'
        return f'<pre class="code-block">[Mermaid Diagram]\n{code}</pre>'

    processed_md = _MERMAID_BLOCK_RE.sub(_mermaid_word_replacer, md_text)

    # Convert Markdown to HTML with Pygments codehilite
    body_html = md.markdown(processed_md, extensions=_MD_EXTENSIONS)
    body_html = body_html.replace('✅', '✓').replace('❌', '✗')

    soup = BeautifulSoup(body_html, 'html.parser')

    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                _set_cjk_font(run)
            continue

        tag = element.name.lower() if element.name else ''

        # Headings (h1 - h6)
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            h = doc.add_heading(level=min(level, 9))
            for child in element.children:
                _append_node_to_paragraph(h, child, theme)

            # Apply theme color to headings
            h_color = theme.docx_colors.h1 if level == 1 else (
                theme.docx_colors.h2 if level == 2 else theme.docx_colors.h3
            )
            for run in h.runs:
                _set_cjk_font(run, size_pt=max(18 - level * 2, 11))
                run.font.color.rgb = h_color
            continue

        # Paragraph
        if tag == 'p':
            p_text = element.get_text().strip()
            if p_text in mermaid_images:
                img_path = mermaid_images[p_text]
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                continue

            p = doc.add_paragraph()
            for child in element.children:
                _append_node_to_paragraph(p, child, theme)
            continue

        # Lists (ul / ol)
        if tag in ('ul', 'ol'):
            list_style = 'List Bullet' if tag == 'ul' else 'List Number'
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style=list_style)
                for child in li.children:
                    if child.name in ('ul', 'ol'):
                        continue
                    _append_node_to_paragraph(p, child, theme)
            continue

        # Blockquote
        if tag == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run('│ ')
            for child in element.children:
                _append_node_to_paragraph(p, child, theme, is_italic=True)
            for run in p.runs:
                run.font.color.rgb = theme.docx_colors.quote
            continue

        # Code block (pre / codehilite)
        if tag in ('pre', 'div') and 'codehilite' in element.get('class', []):
            code_pre = element.find('pre') or element
            p = doc.add_paragraph()
            for child in code_pre.children:
                _append_node_to_paragraph(p, child, theme, is_code=True)

            pPr = p._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), theme.docx_colors.code_block_bg)
            shd.set(qn('w:val'), 'clear')
            pPr.append(shd)
            continue

        if tag == 'pre':
            code_text = element.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            _set_cjk_font(run, font_name='Noto Sans Mono CJK TC', size_pt=9.5)
            run.font.color.rgb = theme.docx_colors.body
            pPr = p._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), theme.docx_colors.code_block_bg)
            shd.set(qn('w:val'), 'clear')
            pPr.append(shd)
            continue

        # Table
        if tag == 'table':
            rows = element.find_all('tr')
            if rows:
                num_cols = max(len(r.find_all(['th', 'td'])) for r in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                for ri, tr in enumerate(rows):
                    cells = tr.find_all(['th', 'td'])
                    for ci, cell_el in enumerate(cells):
                        if ci < num_cols:
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            for child in cell_el.children:
                                _append_node_to_paragraph(p, child, theme, is_bold=(ri == 0))
                            for run in p.runs:
                                _set_cjk_font(run, size_pt=10)

                            if ri == 0:
                                tcPr = cell._element.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:fill'), theme.docx_colors.table_header_bg)
                                shd.set(qn('w:val'), 'clear')
                                tcPr.append(shd)
            continue

        # Horizontal rule
        if tag == 'hr':
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            continue

    tmp_dir = tempfile.mkdtemp(dir=EXPORT_DIR, prefix='word_')
    out_path = os.path.join(tmp_dir, filename)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Gradio UI
# ---------------------------------------------------------------------------

_SAMPLE_MD = '''# Markdown 預覽範例

這是一個 **Markdown 預覽器**，支援 *中文* 與各種格式。

## 功能列表

- ✅ 標題 (h1 ~ h6)
- ✅ **粗體**、*斜體*、`行內程式碼`
- ✅ 清單 (有序 / 無序)
- ✅ 表格
- ✅ 引言區塊
- ✅ 程式碼區塊 (含語法高亮)
- ✅ Mermaid 流程圖

## 表格範例

| 功能 | 狀態 | 備註 |
|------|------|------|
| 中文支援 | ✅ | 完整 CJK 支援 |
| Mermaid | ✅ | 流程圖、序列圖等 |
| 匯出 PDF | ✅ | 支援中文字型 & 主題配色 |
| 匯出 Word | ✅ | .docx 格式 & 語法高亮 |

## 程式碼區塊 (Python)

```python
def hello_world(name: str) -> str:
    # 這是註解說明
    greeting = f'你好，{name}！'
    print(greeting)
    return greeting

hello_world('使用者')
```

> 這是一段引言，支援中文排版與主題配色。

## Mermaid 流程圖

```mermaid
graph TD
    A[開始] --> B{是否有資料?}
    B -->|有| C[處理資料]
    B -->|無| D[等待輸入]
    C --> E[輸出結果]
    D --> B
```

---

*感謝使用本工具！*
'''

_GRADIO_CUSTOM_CSS = '''
.gradio-container {
    max-width: 100% !important;
}
.tabs > .tab-content {
    width: 100% !important;
}
.tabitem {
    width: 100% !important;
}
div[data-testid="html"], .gradio-html {
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
    box-shadow: none !important;
}
.preview-iframe {
    width: 100% !important;
    height: 850px !important;
    min-height: 850px !important;
    border: 1px solid #cbd5e1;
    border-radius: 8px;
}
'''

_DOWNLOAD_JS = '''
(url) => {
    if (url) {
        const rawName = url.split('/').pop();
        const filename = decodeURIComponent(rawName);
        const a = document.createElement('a');
        a.href = url;
        a.download = filename;
        document.body.appendChild(a);
        a.click();
        document.body.removeChild(a);
    }
}
'''


def create_app() -> gr.Blocks:
    '''
    Construct the Gradio UI blocks application with tabs for word counting
    and markdown preview with theme switching and export capabilities.

    Returns:
        gr.Blocks instance with two tabs: word count and markdown preview.
    '''
    theme = gr.themes.Soft(
        primary_hue='indigo',
        secondary_hue='slate',
    )

    with gr.Blocks(
        theme=theme,
        title='文字工具箱 - 字數統計 & Markdown 預覽',
        css=_GRADIO_CUSTOM_CSS,
    ) as demo:
        gr.Markdown(
            '''
            # 📝 文字工具箱
            字數統計 · Markdown 即時預覽 (含多主題/語法高亮) · 匯出 PDF / Word
            '''
        )

        with gr.Tabs():
            # =============== Tab 1: Word Count ===============
            with gr.Tab('📊 字數統計'):
                with gr.Row(equal_height=False):
                    with gr.Column(scale=3, min_width=400):
                        text_input = gr.Textbox(
                            label='請貼上或輸入文字',
                            placeholder='在此處貼上文字...',
                            lines=14,
                            max_lines=30,
                            autofocus=True,
                        )
                        with gr.Row():
                            clear_btn = gr.Button(
                                '清空內容', variant='secondary',
                            )

                    with gr.Column(scale=2, min_width=300):
                        gr.Markdown('### 📊 統計結果')
                        with gr.Row():
                            total_chars_num = gr.Number(
                                label='總字數 (含空格/換行)',
                                value=0, precision=0, interactive=False,
                            )
                            chars_no_spaces_num = gr.Number(
                                label='不含空格/換行字數',
                                value=0, precision=0, interactive=False,
                            )
                        with gr.Row():
                            cjk_count_num = gr.Number(
                                label='中文字數 (漢字)',
                                value=0, precision=0, interactive=False,
                            )
                            cjk_punct_num = gr.Number(
                                label='全形/中文標點',
                                value=0, precision=0, interactive=False,
                            )
                        with gr.Row():
                            english_words_num = gr.Number(
                                label='英文字數 (Words)',
                                value=0, precision=0, interactive=False,
                            )
                            digit_count_num = gr.Number(
                                label='數字個數',
                                value=0, precision=0, interactive=False,
                            )
                        with gr.Row():
                            line_count_num = gr.Number(
                                label='總行數',
                                value=0, precision=0, interactive=False,
                            )

                count_outputs = [
                    total_chars_num, chars_no_spaces_num,
                    cjk_count_num, cjk_punct_num,
                    english_words_num, digit_count_num, line_count_num,
                ]

                text_input.change(
                    fn=analyze_text,
                    inputs=[text_input],
                    outputs=count_outputs,
                )

                clear_btn.click(
                    fn=clear_input,
                    inputs=[],
                    outputs=[text_input] + count_outputs,
                )

            # =============== Tab 2: Markdown Preview ===============
            with gr.Tab('🔍 Markdown 預覽'):
                with gr.Row(equal_height=False):
                    with gr.Column(scale=1, min_width=400):
                        md_input = gr.Textbox(
                            label='Markdown 原始碼',
                            placeholder='在此輸入 Markdown 文字...',
                            lines=26,
                            max_lines=50,
                            value=_SAMPLE_MD,
                        )

                        theme_select = gr.Dropdown(
                            choices=list(THEMES.keys()),
                            value='Light',
                            label='🎨 選擇渲染主題',
                            interactive=True,
                        )

                        with gr.Row():
                            export_pdf_btn = gr.Button(
                                '📄 匯出 PDF', variant='primary',
                            )
                            export_word_btn = gr.Button(
                                '📝 匯出 Word', variant='primary',
                            )

                        # Hidden textboxes to store output file URLs for JS download trigger
                        pdf_url_box = gr.Textbox(visible=False)
                        word_url_box = gr.Textbox(visible=False)

                    with gr.Column(scale=1, min_width=400):
                        gr.Markdown('### 👁️ 即時預覽')
                        preview_html = gr.HTML(
                            value=render_preview(_SAMPLE_MD, 'Light'),
                        )

                # Live preview on text or theme change
                md_input.change(
                    fn=render_preview,
                    inputs=[md_input, theme_select],
                    outputs=[preview_html],
                )

                theme_select.change(
                    fn=render_preview,
                    inputs=[md_input, theme_select],
                    outputs=[preview_html],
                )

                # PDF export (instant single-click download with custom filename via chained JS)
                def _do_export_pdf(md_text: str, theme_name: str) -> str:
                    '''
                    Handle PDF export button click and return file URL for JS download.
                    '''
                    path = export_pdf(md_text, theme_name)
                    if path:
                        return f'/gradio_api/file={path}'
                    return ''

                export_pdf_btn.click(
                    fn=_do_export_pdf,
                    inputs=[md_input, theme_select],
                    outputs=[pdf_url_box],
                ).then(
                    fn=None,
                    inputs=[pdf_url_box],
                    js=_DOWNLOAD_JS,
                )

                # Word export (instant single-click download with custom filename via chained JS)
                def _do_export_word(md_text: str, theme_name: str) -> str:
                    '''
                    Handle Word export button click and return file URL for JS download.
                    '''
                    path = export_word(md_text, theme_name)
                    if path:
                        return f'/gradio_api/file={path}'
                    return ''

                export_word_btn.click(
                    fn=_do_export_word,
                    inputs=[md_input, theme_select],
                    outputs=[word_url_box],
                ).then(
                    fn=None,
                    inputs=[word_url_box],
                    js=_DOWNLOAD_JS,
                )

    return demo


if __name__ == '__main__':
    app = create_app()
    app.launch(
        server_name='0.0.0.0',
        server_port=27860,
        share=False,
        allowed_paths=[EXPORT_DIR, tempfile.gettempdir()],
    )
